import os
import re
import io
import cv2
import zipfile
import numpy as np
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree
from copy import deepcopy
from convertor.common import namespaces, fonts, convert_latex_to_omml


def calculate_avg_height_and_spacing(bounding_boxes, texts, tolerance=0.2):
    heights = []
    spacings = []
    bounding_boxes.sort(key=lambda bbox: bbox[1])

    for i in range(len(bounding_boxes) - 1):
        y1_cur, y2_cur = bounding_boxes[i][1], bounding_boxes[i][3]
        y1_next, y2_next = bounding_boxes[i + 1][1], bounding_boxes[i + 1][3]
        if y1_next <= y1_cur <= y2_next or y1_next <= y2_cur <= y2_next:
            continue
        height = y2_cur - y1_cur
        spacing = y1_next - y2_cur
        if height <50:
            heights.append(height)
        if spacing < 50:
            spacings.append(spacing)

    mean_height = np.mean(heights)
    std_height = np.std(heights)
    mean_spacing = np.mean(spacings)
    std_spacing = np.std(spacings)

    filtered_heights = [h for h in heights if abs(h - mean_height) <= tolerance * std_height]
    filtered_spacings = [s for s in spacings if abs(s - mean_spacing) <= tolerance * std_spacing]

    if len(filtered_heights) > 0:
        avg_height = np.mean(filtered_heights)
    else:
        avg_height = mean_height
        # print('filtered_heights is empty')
    if len(filtered_spacings) > 0:
        avg_spacing = np.mean(filtered_spacings)
    else:
        avg_spacing = mean_spacing
        # print('filtered_spacings is empty')
    return avg_height, avg_spacing


def convert_latex_to_omml_element(text, rPr=None):
    omml_element = OxmlElement('m:oMath')
    try:
        omml_tree = convert_latex_to_omml(text)
        for c in omml_tree.getchildren():
            r_pr = etree.Element(f'{{{namespaces["m"]}}}rPr')
            etree.SubElement(r_pr, f'{{{namespaces["m"]}}}sty',
                             attrib={f'{{{namespaces["m"]}}}val': 'p'})
            c.insert(0, r_pr)
            omml_element.append(c)
            if rPr is not None:
                r = omml_element.getchildren()[-1]
                r.insert(0, rPr)
    except Exception as e:
        print(f'Error: {text}')
        print(e)
    return omml_element


def create_omathpara(omml):
    omml_element = OxmlElement('m:oMathPara')
    omml_para_pr = OxmlElement('m:oMathParaPr')
    jc_element = OxmlElement('m:jc')
    jc_element.set(f'{{{namespaces["m"]}}}val', 'left')
    omml_para_pr.append(jc_element)
    omml_element.append(omml_para_pr)
    omml_element.append(omml)
    return omml_element


class HtmlConvertor:
    def __init__(self):
        self.template = os.path.join(os.path.dirname(__file__), 'template.docx')

    def apply_styles(self, dst):
        font_family = font_size = None
        if self.soup.head is not None:
            for child in self.soup.head.children:
                if child.name == 'style':
                    font_family = re.match(r'.*font-family:(.*?);', child.text)
                    if font_family is not None:
                        font_family = font_family.group(1)
                    font_size = re.match(r'.*font-size:(.*?)pt;', child.text)
                    if font_size is not None:
                        font_size = float(font_size.group(1))
                        if font_size > self.font_size:
                            # print(f'font_size: {font_size} -> {self.font_size}')
                            font_size = self.font_size
                        else:
                            self.font_size = font_size
                        font_size = int(font_size * 2)
        if font_size is None:
            font_size = int(self.font_size * 2)
        output_docx = io.BytesIO()
        with zipfile.ZipFile(self.template, 'r') as docx_zip, zipfile.ZipFile(output_docx, 'a', zipfile.ZIP_DEFLATED) as new_docx_zip:
            for item in docx_zip.infolist():
                if item.filename == 'word/styles.xml':
                    content = docx_zip.read(item.filename).decode('utf-8')
                    if font_family is not None:
                        m = re.search(r'w:ascii="Times New Roman" w:eastAsia="\S+"', content)
                        content = content.replace(m.group(0), f'w:ascii="Times New Roman" w:eastAsia="{font_family}"')
                        # print(f'font_family: {font_family}')
                    if font_size is not None:
                        m = re.search(r'<w:sz\s+w:val="\d+"\s*/>', content)
                        content = content.replace(m.group(0), f'<w:sz w:val="{font_size}"/>')
                        m = re.search(r'<w:szCs\s+w:val="\d+"\s*/>', content)
                        content = content.replace(m.group(0), f'<w:szCs w:val="{font_size}"/>')
                        # print(f'font_size: {font_size}')
                    if self.font_size < 14 and self.line_spacing > 0:
                        m = re.search(r'w:line="\d+"', content)
                        val = self.line_spacing / self.line_height
                        if val < 0.1:
                            val = 240
                            # print(f'单倍行距')
                        elif val < 0.9:
                            val = 360
                            # print(f'1.5倍行距')
                        elif val < 1.5:
                            val = 480
                            # print(f'2倍行距')
                        else:
                            val = 720
                            # print(f'3倍行距')
                        content = content.replace(m.group(0), f'w:line="{val}"')
                    content = content.encode()
                else:
                    content = docx_zip.read(item.filename)
                new_docx_zip.writestr(item, content)
        with open(dst, 'wb') as modified_docx_file:
            modified_docx_file.write(output_docx.getvalue())

    def get_span_style(self, node, rPr=None):
        rPr2 = OxmlElement('w:rPr')
        for k, v in node.attrs.items():
            if k == 'style':
                for style in v.split(';'):
                    if style == '':
                        continue
                    key, value = style.split(':')
                    styles = self.get_style(key, value)
                    for s in styles:
                        rPr2.append(s)
        if rPr is not None:
            for element in rPr:
                skip = False
                for element2 in rPr2:
                    if element.tag == element2.tag:
                        skip = True
                        break
                if not skip:
                    rPr2.append(deepcopy(element))
        return rPr2

    def add_span(self, paragraph, node, rPr=None):
        rPr2 = self.get_span_style(node, rPr)
        self.add_runs(paragraph, node.children, rPr2)

    def add_ruby(self, paragraph, node, rPr=None):
        rt = OxmlElement('w:rt')
        rubyBase = OxmlElement('w:rubyBase')
        for child in node.children:
            rPr2 = None
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = child.text
            if child.name == 'rt':
                for c in child.children:
                    if c.name == 'span':
                        rPr2 = self.get_span_style(c, rPr)
                r.append(rPr2)
                r.append(t)
                rt.append(r)
            else:
                if child.name == 'span':
                    rPr2 = self.get_span_style(child, rPr)
                r.append(rPr2)
                r.append(t)
                rubyBase.append(r)
        ruby = OxmlElement('w:ruby')
        ruby.append(rt)
        ruby.append(rubyBase)
        paragraph._element.append(ruby)

    def add_run(self, paragraph, text, rPr=None, omathpara=False):
        matches = re.finditer(r'\\\(.*?\\\)', text)
        i = 0
        if matches is not None:
            for match in matches:
                if i < match.start():
                    run = paragraph.add_run(text[i:match.start()])
                    if rPr is not None:
                        run._element.insert(0, rPr)
                omml = convert_latex_to_omml_element(match.group()[2:-2], rPr)
                if omathpara and len(match.group()) == len(text):
                    omml = create_omathpara(omml)
                paragraph._element.append(omml)
                i = match.end()
        if i < len(text):
            run = paragraph.add_run(text[i:])
            if rPr is not None:
                run._element.insert(0, rPr)

    def add_runs(self, paragraph, children, rPr=None):
        for c in children:
            if c.name == 'sup' or c.name == 'sub':
                run = paragraph.add_run(c.text)
                if c.name == 'sup':
                    run.font.superscript = True
                else:
                    run.font.subscript = True
                if rPr is not None:
                    run._element.insert(0, rPr)
            elif c.name == 'tab':
                run = paragraph.add_run()
                run.add_tab()
            elif c.name == 'pic':
                self.add_picture(paragraph, c)
            elif c.name == 'span':
                self.add_span(paragraph, c, rPr)
            elif c.name == 'ruby':
                self.add_ruby(paragraph, c, rPr)
            else:
                self.add_run(paragraph, c.text, rPr)

    def add_picture(self, paragraph, node):
        if 'id' not in node.attrs:
            return
        id = node.attrs['id']
        if id not in self.pics:
            return
        box = self.pics[id]
        h = box[3] - box[1]
        lines = h / self.line_height
        line_height = 0
        for font in fonts:
            line_height = font['line_height']
            if self.font_size <= font['pt']:
                break
        img = cv2.imread(f'output/tmp/{id}.png')
        h, w, _ = img.shape
        height = int(lines * line_height)
        width = int(w * height / h)
        run = paragraph.add_run()
        run.add_picture(f'output/tmp/{id}.png', width=width, height=height)

    def get_style(self, key, value):
        if key == 'padding-left':
            assert value.endswith('em')
            left_chars = int(float(value[:-2]) * 100)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{namespaces["w"]}}}leftChars', str(left_chars))
            return [ind]
        elif key == 'text-align':
            jc = OxmlElement('w:jc')
            jc.set(f'{{{namespaces["w"]}}}val', value)
            return [jc]
        elif key == 'text-hanging':
            assert value.endswith('em')
            left_chars = int(float(value[:-2]) * 100)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{namespaces["w"]}}}hangingChars', str(left_chars))
            return [ind]
        elif key == 'text-indent':
            assert value.endswith('em')
            left_chars = int(float(value[:-2]) * 100)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{namespaces["w"]}}}firstLineChars', str(left_chars))
            return [ind]
        elif key == 'font-weight':
            b = OxmlElement('w:b')
            bCs = OxmlElement('w:bCs')
            return [b, bCs]
        elif key == 'font-size':
            sz = OxmlElement('w:sz')
            sz.set(f'{{{namespaces["w"]}}}val', str(int(float(value[:-2]) * 2)))
            szCs = OxmlElement('w:szCs')
            szCs.set(f'{{{namespaces["w"]}}}val', str(int(float(value[:-2]) * 2)))
            return [sz, szCs]
        elif key == 'font-family':
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(f'{{{namespaces["w"]}}}ascii', value)
            rFonts.set(f'{{{namespaces["w"]}}}eastAsia', value)
            return [rFonts]
        elif key == 'color':
            color = OxmlElement('w:color')
            colors = {'black': '000000', 'red': 'FF0000', 'blue': '0000FF', 'green': '00FF00'}
            color.set(f'{{{namespaces["w"]}}}val', colors[value])
            return [color]
        elif key == 'text-decoration':
            if value == 'underline':
                u = OxmlElement('w:u')
                u.set(f'{{{namespaces["w"]}}}val', 'single')
                return [u]
            else:
                assert False, f'Unknown style: {key}'
        else:
            assert False, f'Unknown style: {key}'

    def set_paragraph_format(self, paragraph, node):
        rPr = None
        pPr = OxmlElement('w:pPr')
        tab_count = len([c for c in node.children if c.name == 'tab'])
        if tab_count > 0:
            tabs = OxmlElement('w:tabs')
            pos = 8820 // (tab_count + 1)
            for i in range(tab_count):
                tab = OxmlElement('w:tab')
                tab.set(f'{{{namespaces["w"]}}}val', 'left')
                tab.set(f'{{{namespaces["w"]}}}pos', str(pos * (i + 1)))
                tabs.append(tab)
            pPr.append(tabs)
        for k, v in node.attrs.items():
            if k == 'style':
                for style in v.split(';'):
                    key, value = style.split(':')
                    styles = self.get_style(key, value)
                    if key in ['font-weight', 'font-size', 'font-family', 'color']:
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                        for s in styles:
                            rPr.append(s)
                    else:
                        for s in styles:
                            pPr.append(s)
        paragraph._element.append(pPr)
        return rPr

    def set_table_format(self, table, attrs):
        if 'border' not in attrs:
            return
        tbl = table._element
        tbl_pr = tbl.xpath('w:tblPr')
        if not tbl_pr:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        else:
            tbl_pr = tbl_pr[0]
        tbl_borders = OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border_element = OxmlElement(f'w:{border_type}')
            border_element.set(f'{{{namespaces["w"]}}}val', 'single')
            border_element.set(f'{{{namespaces["w"]}}}sz', '4')
            border_element.set(f'{{{namespaces["w"]}}}color', 'auto')
            tbl_borders.append(border_element)
        tbl_pr.append(tbl_borders)

    def set_paragraph(self, paragraph, node):
        rPr = self.set_paragraph_format(paragraph, node)
        children = list(node.children)
        if len(children) == 1:
            if children[0].name == 'pic':
                self.add_picture(paragraph, children[0])
            else:
                self.add_run(paragraph, node.text, rPr, True)
        else:
            self.add_runs(paragraph, children, rPr)

    def read_ocr(self, path):
        max_count = 0
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                if '<table' in line:
                    continue
                if '<pic' in line:
                    match = re.search(r'<pic id="(\d+)"/>', line)
                    if match is not None:
                        id = match.group(1)
                        match = re.search(r'<box>\((\d+),(\d+)\),\((\d+),(\d+)\)</box>', line)
                        self.pics[id] = [int(match.group(1)), int(match.group(2)), int(match.group(3)), int(match.group(4))]
                else:
                    match = re.search(r'<ref>(.*?)</ref>', line)
                    if match is not None:
                        text = match.group(1)
                        count = len(text)
                        match = re.search(r'<box>\((\d+),(\d+)\),\((\d+),(\d+)\)</box>', line)
                        box = [float(match.group(1)), float(match.group(2)), float(match.group(3)), float(match.group(4))]
                        self.lines.append(box)
                        self.texts[str(box)] = text
                        if '`' not in line and count > max_count:
                            max_count = count
                            # self.line_height = float(match.group(4)) - float(match.group(2))
        if max_count > 80:
            self.font_size = 20
        else:
            self.font_size = fonts[0]['pt']
            for font in fonts[1:]:
                if font['count'] > max_count:
                    self.font_size = font['pt']
                else:
                    break
        self.line_height, self.line_spacing = calculate_avg_height_and_spacing(self.lines, self.texts)

    def add_table(self, doc, table):
        num_rows = len(table.find_all('tr'))
        num_cols = max(len(row.find_all(['th', 'td'])) for row in table.find_all('tr'))
        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_format(word_table, table.attrs)
        for i, row in enumerate(table.find_all('tr')):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                try:
                    if cell.has_attr('rowspan'):
                        span = int(cell['rowspan'])
                        word_table.cell(i, j).merge(word_table.cell(i + span - 1, j))
                    elif cell.has_attr('colspan'):
                        span = int(cell['colspan'])
                        word_table.cell(i, j).merge(word_table.cell(i, j + span - 1))
                    else:
                        paragraph = word_table.cell(i, j).paragraphs[0]
                        if '<p' in cell.text:
                            for child in cell.children:
                                if paragraph is None:
                                    paragraph = word_table.cell(i, j).add_paragraph()
                                if child.name is None:
                                    self.add_run(paragraph, child, omathpara=True)
                                elif child.name == 'p':
                                    self.set_paragraph(paragraph, child)
                                paragraph = None
                        else:
                            self.add_runs(paragraph, cell.children)
                except Exception as e:
                    print(e)

    def convert(self, html, output, ocr=None):
        self.lines = []
        self.texts = {}
        self.pics = {}
        self.font_size = fonts[-1]['pt']
        self.line_spacing = 0
        if ocr is not None:
            self.read_ocr(ocr)
        self.soup = BeautifulSoup(html, 'html.parser')
        self.apply_styles(output)
        doc = Document(output)
        paragraph = doc.paragraphs[-1]
        if self.soup.body is None:
            children = self.soup.children
        else:
            children = self.soup.body.children
        for child in children:
            if child.name == 'p':
                if paragraph is None:
                    paragraph = doc.add_paragraph()
                self.set_paragraph(paragraph, child)
                paragraph = None
            elif child.name == 'table':
                self.add_table(doc, child)

        doc.save(output)


if __name__ == '__main__':
    output = r'C:\Users\Administrator\Desktop\word\output.docx'
    with open(r'output/output.txt', 'r', encoding='utf-8') as f:
        html = f.read().replace('\n', '')
    convertor = HtmlConvertor()
    convertor.convert(html, output, 'output/ocr.txt')
