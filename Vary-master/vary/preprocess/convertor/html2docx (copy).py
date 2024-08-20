import os
import re
import io
import zipfile
import fitz
import shutil
import numpy as np
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from lxml import etree
from copy import deepcopy
from convertor.common import namespaces, fonts, convert_latex_to_omml, docx_save_as
from PIL import Image
from spire.doc import FileFormat as SpireFileFormat
from spire.doc import Document as SpireDocument


def calculate_avg_height_and_spacing(bounding_boxes, tolerance=0.2):
    heights = []
    spacings = []
    bounding_boxes.sort(key=lambda bbox: bbox[1])

    for i in range(len(bounding_boxes) - 1):
        y1_cur, y2_cur = bounding_boxes[i][1], bounding_boxes[i][3]
        y1_next, y2_next = bounding_boxes[i + 1][1], bounding_boxes[i + 1][3]
        if y1_next <= y1_cur <= y2_next or y1_next <= y2_cur <= y2_next:
            continue
        height = y2_cur - y1_cur
        spacing = y1_next - y2_cur
        if height < 50:
            heights.append(height)
        if spacing < 50:
            spacings.append(spacing)

    mean_height = np.mean(heights)
    std_height = np.std(heights)
    mean_spacing = np.mean(spacings)
    std_spacing = np.std(spacings)

    filtered_heights = [h for h in heights if abs(h - mean_height) <= tolerance * std_height]
    filtered_spacings = [s for s in spacings if abs(s - mean_spacing) <= tolerance * std_spacing]

    if len(filtered_heights) > 0:
        avg_height = np.mean(filtered_heights)
    else:
        avg_height = mean_height
        # print('filtered_heights is empty')
    if len(filtered_spacings) > 0:
        avg_spacing = np.mean(filtered_spacings)
    else:
        avg_spacing = mean_spacing
        # print('filtered_spacings is empty')
    return avg_height, avg_spacing


def convert_latex_to_omml_element(text, rPr=None):
    omml_element = OxmlElement('m:oMath')
    try:
        omml_tree = convert_latex_to_omml(text)
        for c in omml_tree.getchildren():
            r_pr = etree.Element(f'{{{namespaces["m"]}}}rPr')
            etree.SubElement(r_pr, f'{{{namespaces["m"]}}}sty',
                             attrib={f'{{{namespaces["m"]}}}val': 'p'})
            c.insert(0, r_pr)
            omml_element.append(c)
            if rPr is not None:
                r = omml_element.getchildren()[-1]
                r.insert(0, rPr)
    except Exception as e:
        print(f'latex error: {text}')
        print(e)
    return omml_element


def create_omathpara(omml):
    omml_element = OxmlElement('m:oMathPara')
    omml_para_pr = OxmlElement('m:oMathParaPr')
    jc_element = OxmlElement('m:jc')
    jc_element.set(f'{{{namespaces["m"]}}}val', 'left')
    omml_para_pr.append(jc_element)
    omml_element.append(omml_para_pr)
    omml_element.append(omml)
    return omml_element


def check_font(font):
    if font not in ['黑体', '微软雅黑', 'MicroftJhengHei', '华文细黑',
                    '华文中宋', '宋体', '等线', '华文宋体', '新宋体', '华文仿宋',
                    '仿宋', '楷体', '华文楷体']:
        # print(font, '->', '黑体')
        font = '黑体'
    return font


class HtmlConvertor:
    def __init__(self, workdir):
        self.template = os.path.join(os.path.dirname(__file__), 'template.docx')
        self.workdir = workdir
        self.font_family = '宋体'
        self.spacing_line = 360
        path = os.path.join(workdir, 'tmp.jpg')
        if os.path.exists(path):
            self.image = Image.open(path)
        self.image_id = 0

    def apply_styles(self, dst):
        font_family = font_size = None
        if self.soup.head is not None:
            for child in self.soup.head.children:
                if child.name == 'style':
                    font_family = re.match(r'.*font-family:(.*?);', child.text)
                    if font_family is not None:
                        font_family = check_font(font_family.group(1))
                        self.font_family = font_family
                    font_size = re.match(r'.*font-size:(.*?)pt;', child.text)
                    if font_size is not None:
                        font_size = float(font_size.group(1))
                        if font_size > self.font_size:
                            # print(f'font_size: {font_size} -> {self.font_size}')
                            font_size = self.font_size
                        else:
                            self.font_size = font_size
                        font_size = int(font_size * 2)
        if font_size is None:
            font_size = int(self.font_size * 2)

        if self.line_spacing > 0:
            val = self.line_spacing / self.line_height
            if self.font_size < 14:
                if val < 0.3:
                    val = 240
                    # print(f'单倍行距')
                elif val < 0.95:
                    val = 360
                    # print(f'1.5倍行距')
                elif val < 1.58:
                    val = 480
                    # print(f'2倍行距')
                elif val < 2.25:
                    val = 600
                    # print(f'2.5倍行距')
                else:
                    val = 720
                    # print(f'3倍行距')
            else:
                if val < 1.58:
                    val = 480
                    # print(f'2倍行距')
                elif val < 2.25:
                    val = 600
                    # print(f'2.5倍行距')
                else:
                    val = 720
                    # print(f'3倍行距')
            self.spacing_line = val
        docx_save_as(self.template, dst, font_family, font_size, self.spacing_line)

    def get_span_style(self, node, rPr=None):
        rPr2 = OxmlElement('w:rPr')
        for k, v in node.attrs.items():
            if k == 'style':
                for style in v.split(';'):
                    if style == '':
                        continue
                    key, value = style.split(':')
                    styles = self.get_style(key, value)
                    for s in styles:
                        rPr2.append(s)
        if rPr is not None:
            for element in rPr:
                skip = False
                for element2 in rPr2:
                    if element.tag == element2.tag:
                        skip = True
                        break
                if not skip:
                    rPr2.append(deepcopy(element))
        return rPr2

    def add_span(self, paragraph, node, rPr=None):
        rPr2 = self.get_span_style(node, rPr)
        self.add_runs(paragraph, node.children, rPr2)

    def add_ruby(self, paragraph, node, rPr=None):
        rt = OxmlElement('w:rt')
        rubyBase = OxmlElement('w:rubyBase')
        for child in node.children:
            rPr2 = None
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = child.text
            if child.name == 'rt':
                for c in child.children:
                    if c.name == 'span':
                        rPr2 = self.get_span_style(c, rPr)
                r.append(rPr2)
                r.append(t)
                rt.append(r)
            else:
                if child.name == 'span':
                    rPr2 = self.get_span_style(child, rPr)
                r.append(rPr2)
                r.append(t)
                rubyBase.append(r)
        ruby = OxmlElement('w:ruby')
        ruby.append(rt)
        ruby.append(rubyBase)
        paragraph._element.append(ruby)

    def add_run(self, paragraph, text, rPr=None, omathpara=False):
        matches = re.finditer(r'\\\(.*?\\\)', text)
        i = 0
        if matches is not None:
            for match in matches:
                if i < match.start():
                    run = paragraph.add_run(text[i:match.start()])
                    if rPr is not None:
                        run._element.insert(0, rPr)
                str = match.group()[2:-2]
                if re.match(r'^[()\s]+$', str):
                    run = paragraph.add_run(str)
                    if rPr is not None:
                        run._element.insert(0, rPr)
                else:
                    omml = convert_latex_to_omml_element(str, rPr)
                    if omathpara and len(match.group()) == len(text):
                        omml = create_omathpara(omml)
                    paragraph._element.append(omml)
                i = match.end()
        if i < len(text):
            run = paragraph.add_run(text[i:])
            if rPr is not None:
                run._element.insert(0, rPr)

    def add_runs(self, paragraph, children, rPr=None):
        for c in children:
            if c.name == 'sup' or c.name == 'sub':
                run = paragraph.add_run(c.text)
                if c.name == 'sup':
                    run.font.superscript = True
                else:
                    run.font.subscript = True
                if rPr is not None:
                    run._element.insert(0, rPr)
            elif c.name == 'tab':
                run = paragraph.add_run()
                run.add_tab()
            elif c.name == 'pic' or c.name == 'box':
                self.add_picture(paragraph, c)
            elif c.name == 'span':
                self.add_span(paragraph, c, rPr)
            elif c.name == 'ruby':
                self.add_ruby(paragraph, c, rPr)
            else:
                self.add_run(paragraph, c.text, rPr)
        if len(list(children)) > 1 and any(c.name == 'pic' for c in children):
            pPr = paragraph._element.find('.//w:pPr', namespaces)
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                paragraph._element.insert(0, pPr)
            textAlignment = pPr.find('.//w:textAlignment', namespaces)
            if textAlignment is None:
                textAlignment = OxmlElement('w:textAlignment')
                textAlignment.set(f'{{{namespaces["w"]}}}val', 'top')
                pPr.append(textAlignment)

    def add_picture(self, paragraph, node):
        if node.name == 'pic':
            if 'id' not in node.attrs:
                return
            id = node.attrs['id']
            if id not in self.pics:
                return
            box = self.pics[id]
            path = os.path.join(self.workdir, f'tmp/{id}.png')
            img = Image.open(path)
            img_width, img_height = img.size
        else:
            match = re.match(r'\((\d+),(\d+)\),\((\d+),(\d+)\)', node.text)
            if match is None:
                return
            box = [int(match.group(1)), int(match.group(2)), int(match.group(3)), int(match.group(4))]
            box2 = [box[0] * self.image.width // 1000, box[1] * self.image.height // 1000,
                    box[2] * self.image.width // 1000, box[3] * self.image.height // 1000]
            id = self.image_id
            self.image_id += 1
            path = os.path.join(self.workdir, f'tmp/{id}.png')
            self.image.crop(box2).save(path)
            img_width = box2[2] - box2[0]
            img_height = box2[3] - box2[1]
        h = box[3] - box[1]
        lines = h / self.line_height
        line_height = 0
        for font in fonts:
            line_height = font['line_height']
            if self.font_size <= font['pt']:
                break
        height = int(lines * line_height)
        width = int(img_width * height / img_height)
        if width > 5652000 * 0.9:
            width = 5652000 * 0.9
            height = int(img_height * width / img_width)
        run = paragraph.add_run()
        run.add_picture(path, width=width, height=height)

        # pPr = paragraph._element.find('.//w:pPr', namespaces)
        # if pPr is None:
        #     pPr = OxmlElement('w:pPr')
        #     paragraph._element.insert(0, pPr)
        # spacing = OxmlElement('w:spacing')
        # spacing.set(f'{{{namespaces["w"]}}}line', str(self.spacing_line))
        # spacing.set(f'{{{namespaces["w"]}}}lineRule', 'auto')
        # pPr.append(spacing)

    def get_style(self, key, value):
        if key == 'padding-left':
            assert value.endswith('em')
            left_chars = int(float(value[:-2]) * 100)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{namespaces["w"]}}}leftChars', str(left_chars))
            return [ind]
        elif key == 'text-align':
            jc = OxmlElement('w:jc')
            jc.set(f'{{{namespaces["w"]}}}val', value)
            return [jc]
        elif key == 'text-hanging':
            assert value.endswith('em')
            left_chars = int(float(value[:-2]) * 100)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{namespaces["w"]}}}hangingChars', str(left_chars))
            return [ind]
        elif key == 'text-indent':
            assert value.endswith('em')
            left_chars = int(float(value[:-2]) * 100)
            ind = OxmlElement('w:ind')
            ind.set(f'{{{namespaces["w"]}}}firstLineChars', str(left_chars))
            return [ind]
        elif key == 'font-weight':
            b = OxmlElement('w:b')
            bCs = OxmlElement('w:bCs')
            return [b, bCs]
        elif key == 'font-size':
            sz = OxmlElement('w:sz')
            sz.set(f'{{{namespaces["w"]}}}val', str(int(float(value[:-2]) * 2)))
            szCs = OxmlElement('w:szCs')
            szCs.set(f'{{{namespaces["w"]}}}val', str(int(float(value[:-2]) * 2)))
            return [sz, szCs]
        elif key == 'font-family':
            rFonts = OxmlElement('w:rFonts')
            value = check_font(value)
            rFonts.set(f'{{{namespaces["w"]}}}ascii', value)
            rFonts.set(f'{{{namespaces["w"]}}}eastAsia', value)
            rFonts.set(f'{{{namespaces["w"]}}}hAnsi', value)
            return [rFonts]
        elif key == 'color':
            color = OxmlElement('w:color')
            colors = {'black': '000000', 'red': 'FF0000', 'blue': '0000FF', 'green': '00FF00'}
            color.set(f'{{{namespaces["w"]}}}val', colors[value])
            return [color]
        elif key == 'text-decoration':
            if value == 'underline':
                u = OxmlElement('w:u')
                u.set(f'{{{namespaces["w"]}}}val', 'single')
                return [u]
            else:
                assert False, f'Unknown style: {key}'
        elif key == 'text-emphasis':
            if value == 'dot':
                em = OxmlElement('w:em')
                em.set(f'{{{namespaces["w"]}}}val', 'dot')
                return [em]
            else:
                assert False, f'Unknown style: {key}'
        else:
            assert False, f'Unknown style: {key}'

    def check_style(self, key, value, text):
        if text is None or text == '':
            return value
        if key not in ['font-size', 'text-align']:
            return value
        pattern = r'[\s()（），。！？；：:\-_\[\]【】{}<>@#$%^&*"`\'.,/\\|~`]+'
        text = re.sub(pattern, '', text)
        box = None
        for t, b in self.boxes.items():
            t = re.sub(pattern, '', t)
            if t == text:
                box = b
                break
        if box is None:
            return value
        if key == 'font-size':
            h = box[3] - box[1]
            font_size = float(value[:-2])
            if font_size > self.font_size:
                if h < self.line_height:
                    return None
            elif h > self.line_height:
                return None
        elif key == 'text-align':
            center1 = (self.left + self.right) / 2
            center2 = (box[0] + box[2]) / 2
            offset = abs(center1 - center2)
            if value == 'center':
                if offset > box[0] - self.left or offset > self.right - box[2]:
                    if box[0] > center1:
                        return 'right'
                    else:
                        return 'left'
            elif value == 'left':
                if box[2] < center1:
                    return 'left'
                elif box[0] > center1:
                    return 'right'
                elif offset < box[0] - self.left and offset < self.right - box[2]:
                    w = (self.right - self.left) * 0.2
                    if box[0] - self.left > w and self.right - box[2] > w:
                        return 'center'
        return value

    def set_paragraph_format(self, paragraph, node):
        rPr = None
        pPr = OxmlElement('w:pPr')
        tab_count = len([c for c in node.children if c.name == 'tab'])
        if tab_count > 0:
            tabs = OxmlElement('w:tabs')
            pos = 8820 // (tab_count + 1)
            for i in range(tab_count):
                tab = OxmlElement('w:tab')
                tab.set(f'{{{namespaces["w"]}}}val', 'left')
                tab.set(f'{{{namespaces["w"]}}}pos', str(pos * (i + 1)))
                tabs.append(tab)
            pPr.append(tabs)
        align = None
        for k, v in node.attrs.items():
            if k == 'style':
                for style in v.split(';'):
                    try:
                        key, value = style.split(':')
                        if key == 'text-align':
                            align = value
                        value = self.check_style(key, value, node.text)
                        if value is None:
                            continue
                        styles = self.get_style(key, value)
                        if key in ['font-weight', 'font-size', 'font-family', 'color']:
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                            for s in styles:
                                rPr.append(s)
                        else:
                            for s in styles:
                                pPr.append(s)
                    except Exception as e:
                        continue
        if align is None:
            align = self.check_style('text-align', 'left', node.text)
            if align is not None and align != 'left':
                jc = OxmlElement('w:jc')
                jc.set(f'{{{namespaces["w"]}}}val', align)
                pPr.append(jc)
        paragraph._element.append(pPr)
        return rPr

    def set_table_format(self, table, attrs):
        if 'border' not in attrs:
            return
        tbl = table._element
        tbl_pr = tbl.xpath('w:tblPr')
        if not tbl_pr:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        else:
            tbl_pr = tbl_pr[0]
        tbl_borders = OxmlElement('w:tblBorders')
        for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border_element = OxmlElement(f'w:{border_type}')
            border_element.set(f'{{{namespaces["w"]}}}val', 'single')
            border_element.set(f'{{{namespaces["w"]}}}sz', '4')
            border_element.set(f'{{{namespaces["w"]}}}color', 'auto')
            tbl_borders.append(border_element)
        tbl_pr.append(tbl_borders)

    def set_paragraph(self, paragraph, node):
        rPr = self.set_paragraph_format(paragraph, node)
        children = list(node.children)
        if len(children) == 1:
            if children[0].name == 'pic' or children[0].name == 'box':
                self.add_picture(paragraph, children[0])
            else:
                self.add_run(paragraph, node.text, rPr, True)
        else:
            self.add_runs(paragraph, children, rPr)

    def set_node(self, node):
        if node.name == 'p':
            paragraph = self.doc.add_paragraph()
            self.set_paragraph(paragraph, node)
        elif node.name == 'table':
            self.add_table(node)
        elif node.name == 'div':
            self.doc.add_section(WD_SECTION.CONTINUOUS)
            for child in node.children:
                if child.name == 'div':
                    for c in child.children:
                        self.set_node(c)
            paragraph = self.doc.paragraphs[-1]
            drawing = paragraph._element.find('.//w:drawing', namespaces)
            self.doc.add_section(WD_SECTION.CONTINUOUS)
            section = self.doc.sections[-2]
            cols = section._sectPr.xpath('./w:cols')[0]
            cols.set(f'{{{namespaces["w"]}}}num', '2')
            del cols.attrib[f'{{{namespaces["w"]}}}sep']
            if drawing is not None:
                extent = drawing.find('.//wp:extent', namespaces)
                cx = int(extent.get('cx'))
                w = int(cx / 914400 / 6.18 * 8500)
                cols.set(f'{{{namespaces["w"]}}}space', '0')
                cols.set(f'{{{namespaces["w"]}}}equalWidth', '0')
                for val in [8500-w, w]:
                    col = OxmlElement('w:col')
                    col.set(f'{{{namespaces["w"]}}}w', str(val))
                    col.set(f'{{{namespaces["w"]}}}space', '0')
                    cols.append(col)

    def read_ocr(self, path):
        max_count = 0
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                if '<table' in line:
                    continue
                if '<pic' in line:
                    match = re.search(r'<pic id="(\d+)"/>', line)
                    if match is not None:
                        id = match.group(1)
                        match = re.search(r'<box>\((\d+),(\d+)\),\((\d+),(\d+)\)</box>', line)
                        self.pics[id] = [int(match.group(1)), int(match.group(2)), int(match.group(3)), int(match.group(4))]
                else:
                    match = re.search(r'<ref\b[^>]*>(.*?)</ref>', line)
                    if match is not None:
                        text = match.group(1)
                        count = 0
                        for c in text:
                            if '\u4e00' <= c <= '\u9fff':
                                count += 1
                            else:
                                count += 0.5
                        match = re.search(r'<box>\((\d+),(\d+)\),\((\d+),(\d+)\)</box>', line)
                        box = [float(match.group(1)), float(match.group(2)), float(match.group(3)), float(match.group(4))]
                        self.lines.append(box)
                        self.boxes[text] = box
                        if '`' not in line and count > max_count:
                            max_count = count
                            # self.line_height = float(match.group(4)) - float(match.group(2))
        if max_count > 60:
            self.font_size = 20
        else:
            self.font_size = fonts[0]['pt']
            for font in fonts[1:]:
                if font['count'] > max_count:
                    self.font_size = font['pt']
                else:
                    break
        self.line_height, self.line_spacing = calculate_avg_height_and_spacing(self.lines)
        self.left = min(box[0] for box in self.lines)
        self.right = max(box[2] for box in self.lines)

    def add_table(self, table):
        num_rows = len(table.find_all('tr'))
        num_cols = max(len(row.find_all(['th', 'td'])) for row in table.find_all('tr'))
        word_table = self.doc.add_table(rows=num_rows, cols=num_cols)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_format(word_table, table.attrs)
        for i, row in enumerate(table.find_all('tr')):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                try:
                    if cell.has_attr('rowspan'):
                        span = int(cell['rowspan'])
                        word_table.cell(i, j).merge(word_table.cell(i + span - 1, j))
                    elif cell.has_attr('colspan'):
                        span = int(cell['colspan'])
                        word_table.cell(i, j).merge(word_table.cell(i, j + span - 1))
                    else:
                        paragraph = word_table.cell(i, j).paragraphs[0]
                        if '<p' in cell.text:
                            for child in cell.children:
                                if paragraph is None:
                                    paragraph = word_table.cell(i, j).add_paragraph()
                                if child.name is None:
                                    self.add_run(paragraph, child, omathpara=True)
                                elif child.name == 'p':
                                    self.set_paragraph(paragraph, child)
                                paragraph = None
                        else:
                            self.add_runs(paragraph, cell.children)
                except Exception as e:
                    print('table error', str(e))

    def check_pic(self, soup):
        ids = []
        for pic in soup.find_all('pic'):
            if 'id' not in pic.attrs:
                continue
            id = pic.attrs['id']
            if id not in self.pics:
                continue
            if id in ids:
                return False
            ids.append(id)
        for id in self.pics.keys():
            w = self.pics[id][2] - self.pics[id][0]
            h = self.pics[id][3] - self.pics[id][1]
            if w < 76 and h < 76:
                continue
            if id not in ids:
                return False
        return True

    def convert(self, html_str, output, ocr=None):
        self.path = output
        html_str = html_str.replace('\n', '')
        html_str = re.sub(r'(<tab/>\s*)+', r'<tab/>', html_str)

        matches = re.finditer(r'\((.*?)\)', html_str)
        for match in matches:
            if re.match(r'^[ _]*$', match.group(1)):
                html_str = html_str[:match.start()+1] + ' '*len(match.group(1)) + html_str[match.end()-1:]
        matches = re.finditer(r'\(\\\)(.*?)\\\(\)', html_str)
        for match in matches:
            if re.match(r'^[ _]*$', match.group(1)):
                html_str = html_str[:match.start()+3] + ' '*len(match.group(1)) + html_str[match.end()-3:]

        # empty_p_pattern = r'(?<!\n)<p(?=\s|>)[^>]*>\s*<\/p>(?!.*\n\s*<p(?=\s|>)[^>]*>\s*<\/p>)'
        # matches = [match.span() for match in re.finditer(empty_p_pattern, html_str)]
        # if len(matches) == 1:
        #     html_str = html_str[:matches[0][0]] + html_str[matches[0][1]:]
        # elif len(matches) > 1:
        #     matches.reverse()
        #     if matches[0][0] != matches[1][1]:
        #         html_str = html_str[:matches[0][0]] + html_str[matches[0][1]:]
        #     for i in range(1, len(matches)-1):
        #         if matches[i][1] != matches[i-1][0] and matches[i][0] != matches[i+1][1]:
        #             html_str = html_str[:matches[i][0]] + html_str[matches[i][1]:]
        #     if matches[-1][1] != matches[-2][0]:
        #         html_str = html_str[:matches[-1][0]] + html_str[matches[-1][1]:]

        matches = re.finditer(r'\\\(.*?\\\)', html_str)
        i = 0
        str = ''
        if matches is not None:
            for match in matches:
                if i < match.start():
                    text = html_str[i:match.start()]
                    text = text.replace('□', f'<span style="font-family:{self.font_family};">□</span>')
                    text = text.replace('○', f'<span style="font-family:{self.font_family};">○</span>')
                    text = text.replace('〇', f'<span style="font-family:{self.font_family};">○</span>')
                    str += text
                text = match.group()
                text = text.replace('○', '〇')
                text = text.replace('★', '\star ')
                str += text
                i = match.end()
        if i < len(html_str):
            text = html_str[i:]
            text = text.replace('□', f'<span style="font-family:{self.font_family};">□</span>')
            text = text.replace('○', f'<span style="font-family:{self.font_family};">○</span>')
            text = text.replace('〇', f'<span style="font-family:{self.font_family};">○</span>')
            str += text
        html_str = str

        self.lines = []
        self.boxes = {}
        self.pics = {}
        self.font_size = fonts[-1]['pt']
        self.line_spacing = 0
        if ocr is not None:
            self.read_ocr(ocr)
        self.soup = BeautifulSoup(html_str, 'html.parser')
        # if not self.check_pic(self.soup):
        #     print(output, 'pic error')
        self.apply_styles(output)
        self.doc = Document(output)
        for child in self.soup.body.children:
            if child.name == 'section':
                for c in child.children:
                    self.set_node(c)
                # self.doc.add_section(WD_SECTION.CONTINUOUS)
                # section = self.doc.sections[-2]
                # cols = section._sectPr.xpath('./w:cols')[0]
                # cols.set(f'{{{namespaces["w"]}}}num', '2')
            else:
                self.set_node(child)
        self.doc.save(output)
        # self.OnePage(output)

    def PageCount(self, path):
        document = SpireDocument()
        document.LoadFromFile(path)
        pdf_path = path[:-5] + '.pdf'
        document.SaveToFile(pdf_path, SpireFileFormat.PDF)
        document.Close()
        doc = fitz.open(str(pdf_path))
        page_count = doc.page_count
        doc.close()
        os.remove(pdf_path)
        return page_count

    def OnePage(self, path):
        tmp_path = path[:-5] + '-tmp.docx'
        page_count = self.PageCount(path)
        if page_count < 2:
            return
        if self.font_size > 12 or self.font_family in ['微软雅黑', 'MicroftJhengHei', '华文细黑', '华文中宋',
                                                       '等线', '华文宋体', '华文仿宋', '华文楷体']:
            for pt in range(self.spacing_line-24, 480-24, -24):
                docx_save_as(path, tmp_path, None, None, pt)
                page_count = self.PageCount(tmp_path)
                if page_count < 2:
                    shutil.move(tmp_path, path)
                    # print(path, self.spacing_line, '->', pt)
                    break
            if page_count > 1:
                self.spacing_line = 480
                if self.font_size > 12:
                    self.font_size = 12
                if self.font_family in ['微软雅黑', 'MicroftJhengHei', '华文细黑', '华文中宋',
                                        '等线', '华文宋体', '华文仿宋', '华文楷体']:
                    self.font_family = {'微软雅黑': '黑体', 'MicroftJhengHei': '黑体', '华文细黑': '黑体', '华文中宋': '宋体',
                                        '等线': '黑体', '华文宋体': '新宋体', '华文仿宋': '仿宋', '华文楷体': '楷体'}[self.font_family]
        if page_count > 1:
            for pt in range(self.spacing_line-24, 240-24, -24):
                docx_save_as(path, tmp_path, self.font_family, self.font_size*2, pt)
                page_count = self.PageCount(tmp_path)
                if page_count < 2:
                    shutil.move(tmp_path, path)
                    # print(path, self.spacing_line, '->', pt)
                    break
        if page_count > 1:
            print(path, 'page count', page_count)
        if os.path.exists(tmp_path):
            os.remove(tmp_path)



if __name__ == '__main__':
    output = r'C:\Users\Administrator\Desktop\word\output.docx'
    with open(r'output/output.txt', 'r', encoding='utf-8') as f:
        html = f.read().replace('\n', '')
    convertor = HtmlConvertor('output')
    convertor.convert(html, output, 'output/ocr.txt')
